"""
Document Text Extractor — EnterpriseMind AI.

Extract text from PDF, DOCX, dan TXT files.

Hybrid Extraction v3 untuk PDF:
- PyMuPDF4LLM sebagai default extractor (cepat, akurat untuk teks)
- page.find_tables() untuk deteksi tabel (akurat, bukan threshold manual)
- Docling hanya untuk halaman dengan tabel
- RapidOCR sebagai fallback ketika Docling gagal
- Hash-based deduplication untuk efisiensi
"""
import hashlib
import logging
import re
from pathlib import Path
from typing import TypedDict

logger = logging.getLogger(__name__)


# ------------------------------------------------------------------ #
# Data Structures
# ------------------------------------------------------------------ #

class PageExtraction(TypedDict):
    """Hasil ekstraksi per halaman PDF."""
    page_number: int
    text: str
    extraction_method: str  # "pymupdf4llm" | "docling" | "ocr" | "ocr_failed" | "docling_failed"
    source_file: str
    char_count: int


# ------------------------------------------------------------------ #
# Hash-based Deduplication Helpers
# ------------------------------------------------------------------ #

def normalize_for_hash(text: str) -> str:
    """Normalisasi teks untuk hash: lowercase + hilangkan whitespace berlebih."""
    return " ".join(text.lower().split())


def content_hash(text: str) -> str:
    """Generate SHA-256 hash dari teks yang sudah dinormalisasi."""
    normalized = normalize_for_hash(text)
    return hashlib.sha256(normalized.encode()).hexdigest()


# ------------------------------------------------------------------ #
# Public API
# ------------------------------------------------------------------ #

def extract_text(file_path: str, file_type: str) -> str:
    """Extract text from document. Raises FileNotFoundError, ValueError, RuntimeError."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File tidak ditemukan: {file_path}")

    file_type = file_type.lower().strip(".")

    supported_types = {"pdf", "docx", "txt"}
    if file_type not in supported_types:
        raise ValueError(
            f"Tipe file '{file_type}' tidak didukung. "
            f"Tipe yang didukung: {supported_types}"
        )

    logger.info("Mengekstrak teks: file=%s, type=%s", path.name, file_type)

    try:
        if file_type == "txt":
            return _extract_txt(path)
        elif file_type == "pdf":
            return _extract_pdf(path)
        elif file_type == "docx":
            return _extract_docx(path)
        else:
            raise ValueError(f"Handler untuk '{file_type}' belum diimplementasi")

    except (FileNotFoundError, ValueError):
        raise
    except Exception as e:
        raise RuntimeError(
            f"Gagal mengekstrak teks dari {path.name}: {e}"
        ) from e


def extract_text_with_pages(file_path: str, file_type: str) -> list[PageExtraction] | str:
    """
    Extract text dengan page-level info untuk PDF, atau string biasa untuk DOCX/TXT.
    
    Returns:
        - PDF: List[PageExtraction] dengan metadata per halaman
        - DOCX/TXT: str (backward compatibility)
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File tidak ditemukan: {file_path}")

    file_type = file_type.lower().strip(".")

    if file_type == "pdf":
        return _extract_pdf_with_pages(path)
    else:
        return extract_text(file_path, file_type)


def flatten_pages(pages: list[PageExtraction]) -> str:
    """
    Flatten List[PageExtraction] menjadi string tunggal.
    
    Menjaga urutan halaman dan menambahkan separator untuk kompatibilitas
    dengan modul lain yang mengharapkan string panjang utuh.
    """
    if not pages:
        return ""
    
    parts = []
    for page in pages:
        if page["text"] and page["text"].strip():
            parts.append(page["text"].strip())
    
    return "\n\n---\n\n".join(parts)


# ------------------------------------------------------------------ #
# Private Extractors
# ------------------------------------------------------------------ #

def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _extract_pdf(path: Path) -> str:
    """
    Hybrid PDF extraction: PyMuPDF4LLM + Docling + OCR fallback.
    Mengembalikan string tunggal (backward compatibility).
    """
    pages = _extract_pdf_with_pages(path)
    return flatten_pages(pages)


def _extract_pdf_with_pages(path: Path) -> list[PageExtraction]:
    """
    Hybrid extraction v3: PyMuPDF4LLM + Docling + OCR fallback.
    
    Alur:
    1. PyMuPDF4LLM → extract semua halaman (default, cepat, akurat untuk teks)
    2. page.find_tables() → deteksi halaman dengan tabel (akurat)
    3. Halaman tabel → Docling (do_table_structure=True)
    4. Jika Docling return <!-- image --> → RapidOCR fallback
    5. Gabungkan hasil PyMuPDF4LLM + Docling + OCR
    """
    import pymupdf4llm

    logger.info("Hybrid extraction v3: %s", path.name)

    # Step 1: PyMuPDF4LLM untuk semua halaman
    logger.info("[Extractor] Step 1: PyMuPDF4LLM extraction...")
    md_pages = pymupdf4llm.to_markdown(
        str(path),
        page_chunks=True,
        write_images=False,
    )
    logger.info("[Extractor] PyMuPDF4LLM: %d pages extracted", len(md_pages))

    # Step 2: Deteksi tabel dengan page.find_tables()
    logger.info("[Extractor] Step 2: Detecting tables...")
    table_pages = _detect_table_pages(path)
    logger.info("[Extractor] Tables found: %d pages %s", len(table_pages), table_pages)

    # Step 3: Docling untuk halaman tabel
    docling_results: dict[int, PageExtraction] = {}
    if table_pages:
        logger.info("[Extractor] Step 3: Docling extraction for %d table pages...", len(table_pages))
        docling_results = _extract_with_docling_batched(
            path, table_pages, do_table_structure=True
        )

    # Step 4: OCR fallback untuk halaman yang Docling gagal
    failed_pages = []
    for p in table_pages:
        if p not in docling_results:
            failed_pages.append(p)
        elif re.search(r'<!-- (image|Start of picture)', docling_results[p]["text"]):
            failed_pages.append(p)

    ocr_results: dict[int, PageExtraction] = {}
    if failed_pages:
        logger.info("[Extractor] Step 4: OCR fallback for %d failed pages...", len(failed_pages))
        ocr_results = _ocr_fallback(path, failed_pages)

    # Step 5: Gabungkan hasil
    logger.info("[Extractor] Step 5: Merging results...")
    pages: list[PageExtraction] = []

    for i, page_data in enumerate(md_pages):
        if i in ocr_results:
            # OCR results (Docling gagal)
            pages.append(ocr_results[i])
        elif i in docling_results:
            # Docling results (tabel berhasil extract)
            pages.append(docling_results[i])
        else:
            # PyMuPDF4LLM results (default)
            text = page_data.get("text", "").strip()
            text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL).strip()
            text = strip_watermarks(text)
            text = clean_extraction_text(text)
            pages.append(PageExtraction(
                page_number=i + 1,
                text=text,
                extraction_method="pymupdf4llm",
                source_file=path.name,
                char_count=len(text),
            ))

    # Statistik
    methods = {}
    for p in pages:
        m = p["extraction_method"]
        methods[m] = methods.get(m, 0) + 1
    logger.info("[Extractor] Final: %s", methods)

    return pages


# ------------------------------------------------------------------ #
# Table Detection (PyMuPDF find_tables)
# ------------------------------------------------------------------ #

def _detect_table_pages(path: Path) -> list[int]:
    """
    Deteksi halaman yang memiliki tabel menggunakan PyMuPDF find_tables().
    Mengembalikan list 0-indexed page indices.
    """
    import pymupdf

    table_pages = []
    doc = pymupdf.open(str(path))

    for i in range(len(doc)):
        try:
            tables = doc[i].find_tables()
            if tables.tables:
                table_pages.append(i)
        except Exception as e:
            logger.warning("[TableDetector] Error detecting tables on page %d: %s", i + 1, e)

    doc.close()

    return table_pages


# ------------------------------------------------------------------ #
# OCR Fallback (RapidOCR)
# ------------------------------------------------------------------ #

def _ocr_fallback(path: Path, page_numbers: list[int]) -> dict[int, PageExtraction]:
    """
    OCR fallback untuk halaman yang Docling gagal extract.
    Menggunakan RapidOCR (lebih cepat dari EasyOCR).
    
    Args:
        path: Path ke PDF file
        page_numbers: List 0-indexed page indices yang perlu OCR
    
    Returns:
        Dict mapping page_index → PageExtraction
    """
    from rapidocr_onnxruntime import RapidOCR
    import fitz  # PyMuPDF

    ocr = RapidOCR()
    results: dict[int, PageExtraction] = {}

    for page_idx in page_numbers:
        try:
            # Render halaman ke gambar (150 DPI untuk hemat memory)
            doc = fitz.open(str(path))
            page = doc[page_idx]
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            doc.close()

            # OCR
            result, _ = ocr(img_bytes)

            if result:
                # RapidOCR mengembalikan list [bbox, text, confidence]
                text = "\n".join([line[1] for line in result])
            else:
                text = ""

            text = strip_watermarks(text)
            text = clean_extraction_text(text)

            results[page_idx] = PageExtraction(
                page_number=page_idx + 1,
                text=text,
                extraction_method="ocr",
                source_file=path.name,
                char_count=len(text),
            )

            logger.debug("[OCR] Page %d: %d chars extracted", page_idx + 1, len(text))

        except Exception as e:
            logger.error("[OCR] Failed to extract page %d: %s", page_idx + 1, e)
            results[page_idx] = PageExtraction(
                page_number=page_idx + 1,
                text="",
                extraction_method="ocr_failed",
                source_file=path.name,
                char_count=0,
            )

    return results


# ------------------------------------------------------------------ #
# Watermark Stripping
# ------------------------------------------------------------------ #

WATERMARK_PATTERNS = [
    r'https?://\S+',           # URLs
    r'www\.\S+',               # www URLs
    r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b',  # IP addresses
    r'©\s*\d{4}',             # Copyright notices
    r'Confidential|Draft|Sample|Template',  # Common watermarks
]

def strip_watermarks(text: str) -> str:
    """Hapus watermark umum dari teks."""
    for pattern in WATERMARK_PATTERNS:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # Hapus baris yang hanya berisi whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()


def clean_extraction_text(text: str) -> str:
    """Bersihkan teks hasil extraction: HTML tags, concatenated words, artifacts."""
    if not text:
        return text
    
    # 1. Convert <br> ke line break
    text = re.sub(r'<br\s*/?>', '\n', text)
    
    # 2. Remove semua HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # 3. Decode common HTML entities
    text = text.replace('&reg;', '®').replace('&copy;', '©').replace('&trade;', '™')
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('&nbsp;', ' ').replace('&quot;', '"')
    
    # 4. Fix common OCR concatenations (kata tergabung dari OCR)
    OCR_WORD_FIXES = {
        # Active Listening - 7 Key Skills
        "Beattentive": "Be attentive",
        "Askopen-ended": "Ask open-ended",
        "Askprobing": "Ask probing",
        "Requestclarification": "Request clarification",
        "Beattuned": "Be attuned",
        "reflectfeelings": "reflect feelings",
        "toand": "to and",
        # TNA - Common concatenations
        "Prosesanalisis": "Proses analisis",
        "menetapkankompetensi": "menetapkan kompetensi",
        "organisasikebutuhan": "organisasi kebutuhan",
        "pembelajaranorganisasi": "pembelajaran organisasi",
        "uraianjabatan": "uraian jabatan",
        "kompetensidankebutuhan": "kompetensi dan kebutuhan",
        "pembelajaranperjabatan": "pembelajaran per jabatan",
        "BusinessIssue": "Business Issue",
        "PerformanceIssue": "Performance Issue",
        "CompetencyIssue": "Competency Issue",
    }
    for old, new in OCR_WORD_FIXES.items():
        text = text.replace(old, new)
    
    # 4b. Add space sebelum uppercase setelah lowercase (camelCase split)
    #    "Prosesanalisis" → "Proses analisis" (jika belum di-fix di atas)
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    
    # 5. Add space antara huruf dan digit
    text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
    text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)
    
    # 6. Remove ®, ™, © yang tersisa
    text = re.sub(r'[®™©]', '', text)
    
    # 7. Remove empty lines berlebih
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


# ------------------------------------------------------------------ #
# Mini PDF Builder
# ------------------------------------------------------------------ #

def _build_mini_pdf(path: Path, page_numbers: list[int]) -> bytes | None:
    """
    Buat mini PDF di memori (RAM) yang hanya berisi halaman yang diminta.

    Menggunakan PyMuPDF (fitz) untuk menyalin halaman spesifik dari PDF asli
    ke PDF baru di memori. Ini mengurangi payload yang dikirim ke Docling
    secara drastis.

    Returns:
        bytes dari mini PDF, atau None jika gagal.
    """
    import fitz  # PyMuPDF

    try:
        src_doc = fitz.open(str(path))
        mini_doc = fitz.open()
        for pg in page_numbers:
            mini_doc.insert_pdf(src_doc, from_page=pg, to_page=pg)

        pdf_bytes = mini_doc.tobytes()
        mini_doc.close()
        src_doc.close()

        logger.debug(
            "[Docling] Mini PDF created: %d pages -> %d bytes (%.1f KB)",
            len(page_numbers), len(pdf_bytes), len(pdf_bytes) / 1024,
        )
        return pdf_bytes
    except Exception as e:
        logger.error("[Docling] Gagal membuat mini PDF untuk pages %s: %s", page_numbers, e)
        return None


# ------------------------------------------------------------------ #
# Docling Extraction
# ------------------------------------------------------------------ #

def _extract_with_docling_single_batch(
    path: Path,
    page_numbers: list[int],
    do_table_structure: bool,
) -> dict[int, PageExtraction]:
    """Extract satu batch halaman dengan Docling menggunakan mini PDF."""
    import json
    import httpx
    from app.core.config import settings

    docling_url = settings.DOCLING_URL
    results: dict[int, PageExtraction] = {}

    # Buat mini PDF di memori (hanya halaman yang diminta)
    mini_pdf_bytes = _build_mini_pdf(path, page_numbers)
    if mini_pdf_bytes is None:
        logger.error("[Docling] Skip batch %s: gagal membuat mini PDF", page_numbers)
        return results

    # page_range TIDAK diperlukan lagi karena mini PDF sudah hanya berisi halaman target
    options = {
        "to_formats": ["md"],
        "do_ocr": False,
        "ocr": False,
        "do_table_structure": do_table_structure,
    }

    try:
        response = httpx.post(
            f"{docling_url}/v1/convert/file",
            files={"files": ("mini_batch.pdf", mini_pdf_bytes, "application/pdf")},
            data={"options": json.dumps(options), "do_ocr": "false"},
            timeout=1200.0,  # 20 menit per batch
        )
        response.raise_for_status()
        result = response.json()

        document = result.get("document", {})
        markdown_text = document.get("md_content", "")

        if markdown_text:
            # Strip watermarks dari hasil docling
            markdown_text = strip_watermarks(markdown_text)
            markdown_text = clean_extraction_text(markdown_text)
            page_texts = _split_docling_output(markdown_text, page_numbers)

            for page_idx, text in zip(page_numbers, page_texts):
                results[page_idx] = PageExtraction(
                    page_number=page_idx + 1,
                    text=text.strip(),
                    extraction_method="docling",
                    source_file=path.name,
                    char_count=len(text.strip()),
                )

    except Exception as e:
        logger.error("[Docling] Batch extraction failed for pages %s: %s", page_numbers, e)

    return results


def _extract_with_docling_batched(
    path: Path,
    page_numbers: list[int],
    do_table_structure: bool,
) -> dict[int, PageExtraction]:
    """
    Extract halaman dengan Docling menggunakan sub-batching + mini PDF.
    Sleep 2 detik antar batch untuk mencegah Docling kebanjiran request.
    """
    import time

    results: dict[int, PageExtraction] = {}
    DOCLING_BATCH_SIZE = 2
    DOCLING_BATCH_SLEEP = 2.0  # detik antar batch

    batches = [
        page_numbers[i:i + DOCLING_BATCH_SIZE]
        for i in range(0, len(page_numbers), DOCLING_BATCH_SIZE)
    ]

    logger.info(
        "[Docling] Sub-batching: %d pages -> %d batches of %d (do_table_structure=%s, sleep=%.0fs)",
        len(page_numbers), len(batches), DOCLING_BATCH_SIZE, do_table_structure, DOCLING_BATCH_SLEEP
    )

    for batch_idx, batch_pages in enumerate(batches):
        logger.info(
            "[Docling] Batch %d/%d: pages %s",
            batch_idx + 1, len(batches), batch_pages
        )
        batch_results = _extract_with_docling_single_batch(
            path, batch_pages, do_table_structure
        )
        results.update(batch_results)

        # Jeda antar batch (kecuali batch terakhir)
        if batch_idx < len(batches) - 1:
            logger.debug("[Docling] Sleep %.0fs sebelum batch berikutnya...", DOCLING_BATCH_SLEEP)
            time.sleep(DOCLING_BATCH_SLEEP)

    return results


def _split_docling_output(markdown_text: str, page_numbers: list[int]) -> list[str]:
    """
    Split Docling output menjadi per-halaman.
    Docling biasanya menambahkan page break marker.
    """
    # Coba split dengan page break marker yang umum
    page_pattern = r'(?:---\s*page\s+\d+\s*---|\[Page\s+\d+\]|\f)'
    parts = re.split(page_pattern, markdown_text, flags=re.IGNORECASE)
    
    # Filter empty strings
    parts = [p.strip() for p in parts if p.strip()]
    
    # Jika jumlah parts cocok dengan page_numbers, gunakan langsung
    if len(parts) == len(page_numbers):
        return parts
    
    # Fallback: bagi rata berdasarkan karakter
    if not parts:
        return [""] * len(page_numbers)
    
    # Distribusi teks ke halaman yang diminta
    total_len = len(markdown_text)
    per_page = total_len // len(page_numbers) if page_numbers else total_len
    
    results = []
    for i in range(len(page_numbers)):
        start = i * per_page
        end = start + per_page if i < len(page_numbers) - 1 else total_len
        results.append(markdown_text[start:end])
    
    return results


# ------------------------------------------------------------------ #
# DOCX & TXT Extractors
# ------------------------------------------------------------------ #

def _extract_docx(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    return "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def detect_file_type(filename: str) -> str:
    """Detect file type from extension. Raises ValueError if unsupported."""
    ext = Path(filename).suffix.lower().strip(".")
    supported = {"pdf", "docx", "txt"}
    if ext not in supported:
        raise ValueError(
            f"Ekstensi '.{ext}' tidak didukung. "
            f"Format yang didukung: {', '.join(f'.{s}' for s in supported)}"
        )
    return ext
